import base64
import io
import pandas as pd
import numpy as np
import tempfile
import re

import dash
from dash import dcc, html, Input, Output, State, no_update
import dash_bootstrap_components as dbc
import plotly.graph_objects as go

from data_parser import load_and_validate_data

# --- Report Generation Imports ---
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches

# --- Mapbox Configuration ---
# Using a free mapbox token - replace with your own for production
mapbox_access_token = "pk.eyJ1IjoicGxvdGx5bWFwYm94IiwiYSI6ImNrOWJqb2F4djBnMjEzbG50amg0dnJieG4ifQ.Zme1-UoBQEiCRLbXP_-H3w"

# =============================================================================
# App Initialization
# =============================================================================
app = dash.Dash(__name__, external_stylesheets=[dbc.themes.CYBORG], suppress_callback_exceptions=True)
server = app.server

def card(children, **kwargs):
    """A helper function to create a Bootstrap card component."""
    return dbc.Card(dbc.CardBody(children), **kwargs)

# =============================================================================
# Application Layout
# =============================================================================
app.layout = dbc.Container(fluid=True, children=[
    dcc.Store(id='dataframe-store'),
    dcc.Interval(id='sim-interval', interval=200, disabled=True), # Drives the simulation

    html.H1("Aerospace Flight Data Analyzer & Simulator", className="text-center my-4 text-primary"),

    dbc.Row([
        # --- Left Column: Controls ---
        dbc.Col(width=4, children=[
            card([
                html.H4("Control Panel", className="card-title"),
                dcc.Upload(
                    id='upload-data',
                    children=html.Div(['Drag and Drop or ', html.A('Select Flight Data File')]),
                    style={
                        'width': '100%', 'height': '60px', 'lineHeight': '60px',
                        'borderWidth': '1px', 'borderStyle': 'dashed', 'borderRadius': '5px',
                        'textAlign': 'center', 'margin': '10px 0'
                    }
                ),
                html.Div(id='upload-status'),
                html.Div([
                    html.Button("Download PDF Report", id="download-pdf-btn", className="btn btn-primary mt-2 me-2"),
                    dcc.Download(id="download-pdf"),
                    html.Button("Download Word Report", id="download-word-btn", className="btn btn-secondary mt-2"),
                    dcc.Download(id="download-word")
                ], id='download-buttons', style={'display': 'none'}) # Hidden until data is loaded
            ]),
            # --- Dynamically populated cards ---
            card([], id='parser-report-card', className="mt-4", style={'display': 'none'}),
            card([], id='summary-report-card', className="mt-4", style={'display': 'none'}),
            card([], id='plot-controls-card', className="mt-4", style={'display': 'none'}),
            card([], id='comm-outage-card', className="mt-4", style={'display': 'none'}),
            card([], id='sim-controls-card', className="mt-4", style={'display': 'none'}), # NEW
        ]),

        # --- Right Column: Visualizations ---
        dbc.Col(width=8, children=[
            card(
                dbc.Tabs(id="tabs", active_tab="tab-map", children=[
                    dbc.Tab(dcc.Graph(id='map-graph', style={'height': '80vh'}), label="Flight Path Analysis", tab_id="tab-map"),
                    dbc.Tab(id='sim-tab', label="Flight Simulator", tab_id="tab-sim", children=[ # NEW
                        dcc.Graph(id='sim-graph', style={'height': '70vh'}),
                        html.Div(id='sim-telemetry-display', className='mt-2 p-2 border rounded')
                    ]),
                    dbc.Tab(dcc.Graph(id='main-graph', style={'height': '80vh'}), label="Telemetry Plots", tab_id="tab-2d"),
                    dbc.Tab(dcc.Graph(id='scatter-3d-graph', style={'height': '80vh'}), label="3D Scatter Plot", tab_id="tab-3d"),
                ])
            )
        ])
    ])
])

# =============================================================================
# Helper Functions
# =============================================================================
def generate_summary_report(df, outage_threshold_db=3.0):
    """Analyzes the dataframe to produce a summary report card."""
    comm_cols = [col for col in df.columns if col.startswith('COMM_') and col.endswith('_dB')]
    link_rows = []
    if comm_cols:
        for link in comm_cols:
            link_name = re.search('COMM_(.*)_dB', link).group(1).replace('_', ' ')
            outages = df[link] < outage_threshold_db
            outage_duration_sec = df['Timestamp'][outages].diff().dt.total_seconds().sum()
            link_rows.append(html.Tr([
                html.Td(link_name),
                html.Td(f"{df[link].mean():.2f} dB"),
                html.Td(f"{df[link].min():.2f} dB"),
                html.Td(f"{outage_duration_sec:.1f} sec")
            ]))
    comm_table = dbc.Table(
        [html.Thead(html.Tr([html.Th("Link"), html.Th("Avg Margin"), html.Th("Min Margin"), html.Th("Outage Time")])),
         html.Tbody(link_rows)],
        bordered=True, striped=True, hover=True
    )
    return [html.H4("Flight Summary Report", className="card-title"), comm_table]

def create_report_map_figure(df, comm_link, threshold):
    """Generates the Plotly figure for outage map, used in reports."""
    fig = go.Figure()
    
    # Calculate map center and bounds
    center_lat = df['POS_Latitude_deg'].mean()
    center_lon = df['POS_Longitude_deg'].mean()
    lat_range = df['POS_Latitude_deg'].max() - df['POS_Latitude_deg'].min()
    lon_range = df['POS_Longitude_deg'].max() - df['POS_Longitude_deg'].min()
    
    if comm_link and threshold is not None and comm_link in df.columns:
        # Color code based on communication link status
        df_copy = df.copy()
        df_copy['is_outage'] = df_copy[comm_link] < threshold
        df_copy['segment'] = (df_copy['is_outage'].diff().ne(0)).cumsum()
        
        for _, group in df_copy.groupby('segment'):
            color = '#DC3545' if group['is_outage'].iloc[0] else '#1F77B4'
            name = 'Link Outage' if group['is_outage'].iloc[0] else 'Link OK'
            fig.add_trace(go.Scatter(
                x=group['POS_Longitude_deg'], 
                y=group['POS_Latitude_deg'],
                mode="lines", 
                line=dict(width=3, color=color), 
                name=name,
                hovertemplate='<b>%{fullData.name}</b><br>' +
                             'Lat: %{y:.4f}°<br>' +
                             'Lon: %{x:.4f}°<br>' +
                             '<extra></extra>'
            ))
    else:
        # Simple flight path
        fig.add_trace(go.Scatter(
            x=df['POS_Longitude_deg'], 
            y=df['POS_Latitude_deg'],
            mode="lines+markers", 
            line=dict(width=3, color='#FF851B'),
            marker=dict(size=4, color='#FF851B'),
            name="Flight Path",
            hovertemplate='<b>Flight Path</b><br>' +
                         'Lat: %{y:.4f}°<br>' +
                         'Lon: %{x:.4f}°<br>' +
                         '<extra></extra>'
        ))

    # Add start and end markers
    fig.add_trace(go.Scatter(
        x=[df['POS_Longitude_deg'].iloc[0]], 
        y=[df['POS_Latitude_deg'].iloc[0]],
        mode="markers",
        marker=dict(size=12, color='green', symbol='circle'),
        name="Start",
        hovertemplate='<b>START</b><br>' +
                     'Lat: %{y:.4f}°<br>' +
                     'Lon: %{x:.4f}°<br>' +
                     '<extra></extra>'
    ))
    
    fig.add_trace(go.Scatter(
        x=[df['POS_Longitude_deg'].iloc[-1]], 
        y=[df['POS_Latitude_deg'].iloc[-1]],
        mode="markers",
        marker=dict(size=12, color='red', symbol='square'),
        name="End",
        hovertemplate='<b>END</b><br>' +
                     'Lat: %{y:.4f}°<br>' +
                     'Lon: %{x:.4f}°<br>' +
                     '<extra></extra>'
    ))

    fig.update_layout(
        title="Flight Path Analysis",
        template="plotly_dark",
        xaxis=dict(
            title="Longitude (degrees)",
            range=[center_lon - lon_range*0.6, center_lon + lon_range*0.6]
        ),
        yaxis=dict(
            title="Latitude (degrees)", 
            range=[center_lat - lat_range*0.6, center_lat + lat_range*0.6],
            scaleanchor="x",
            scaleratio=1
        ),
        margin={"r":20,"t":60,"l":60,"b":40},
        legend=dict(
            yanchor="top", 
            y=0.99, 
            xanchor="left", 
            x=0.01,
            bgcolor="rgba(0,0,0,0.5)"
        ),
        hovermode='closest'
    )
    
    return fig

# =============================================================================
# Core Callbacks
# =============================================================================
@app.callback(
    [Output('dataframe-store', 'data'), Output('upload-status', 'children'),
     Output('download-buttons', 'style'),
     Output('parser-report-card', 'children'), Output('parser-report-card', 'style'),
     Output('summary-report-card', 'children'), Output('summary-report-card', 'style'),
     Output('plot-controls-card', 'children'), Output('plot-controls-card', 'style'),
     Output('comm-outage-card', 'children'), Output('comm-outage-card', 'style'),
     Output('sim-controls-card', 'children'), Output('sim-controls-card', 'style')],
    Input('upload-data', 'contents'), State('upload-data', 'filename')
)
def process_upload_and_update_ui(contents, filename):
    """Master callback to handle file upload and initialize the entire UI."""
    if contents is None:
        return [no_update] * 13

    content_type, content_string = contents.split(',')
    decoded = base64.b64decode(content_string)
    try:
        df, report = load_and_validate_data(io.StringIO(decoded.decode('utf-8')))
        if df is None or df.empty:
            raise ValueError("Data parsing failed. Check file format.")

        status = html.Div(f"Loaded: {filename}", className="text-success")
        display_block = {'display': 'block', 'marginTop': '1rem'}
        
        # --- Build UI components ---
        report_children = [html.H4("Data Validation", className="card-title")] + [html.P(w) for w in report['warnings']]
        summary_children = generate_summary_report(df)
        
        plot_controls = [
            html.H4("Telemetry Plotting", className="card-title"), 
            html.Label("Select Subsystems:"),
            dcc.Checklist(
                id='subsystem-checklist',
                options=[{'label': s, 'value': s} for s in report.get('subsystems_found', [])],
                labelStyle={'display': 'block'}
            ),
            # Add 3D scatter plot controls here
            html.Hr(),
            html.H5("3D Scatter Plot Controls"),
            html.Label("X-Axis:"), 
            dcc.Dropdown(
                id='x-axis-selector', 
                options=[{'label': col, 'value': col} for col in df.select_dtypes(include=np.number).columns],
                value='POS_Longitude_deg'
            ),
            html.Label("Y-Axis:", className="mt-2"), 
            dcc.Dropdown(
                id='y-axis-selector',
                options=[{'label': col, 'value': col} for col in df.select_dtypes(include=np.number).columns],
                value='POS_Latitude_deg'
            ),
            html.Label("Z-Axis:", className="mt-2"), 
            dcc.Dropdown(
                id='z-axis-selector',
                options=[{'label': col, 'value': col} for col in df.select_dtypes(include=np.number).columns],
                value='POS_Altitude_ft'
            ),
            html.Label("Color By:", className="mt-2"), 
            dcc.Dropdown(
                id='color-selector',
                options=[{'label': col, 'value': col} for col in df.select_dtypes(include=np.number).columns],
                value=df.select_dtypes(include=np.number).columns[0]
            )
        ]

        comm_link_cols = [col for col in df.columns if col.startswith('COMM_') and col.endswith('_dB')]
        comm_outage_controls = [
            html.H4("Comm Link Analysis", className="card-title"), html.Label("Analyze comms link:"),
            dcc.Dropdown(id='comm-link-selector', options=comm_link_cols, placeholder="Select a link..."),
            html.Label("Outage Threshold (dB):", className="mt-2"),
            dcc.Input(id='outage-threshold-input', type='number', value=3, step=0.5, className="form-control")
        ]
        
        # --- NEW: Simulation Controls ---
        sim_controls = [
            html.H4("Simulation Controls", className="card-title"),
            html.Div([
                dbc.Button("▶ Play", id="play-pause-btn", color="success", className="me-2"),
                dbc.Button("■ Stop", id="stop-btn", color="danger"),
            ], className="d-flex"),
            dcc.Slider(id='time-slider', min=0, max=len(df)-1, value=0, step=1, marks=None, tooltip={"placement": "bottom", "always_visible": True}),
            html.Label("Playback Speed:", className="mt-2"),
            dcc.Dropdown(id='speed-selector', options=[
                {'label': '1x', 'value': 1}, {'label': '10x', 'value': 10},
                {'label': '50x', 'value': 50}, {'label': '100x', 'value': 100},
                {'label': '500x', 'value': 500}], value=100)
        ]

        return (
            df.to_json(orient='split'), status, {'display': 'block'},
            report_children, display_block, summary_children, display_block,
            plot_controls, display_block, comm_outage_controls, display_block,
            sim_controls, display_block
        )

    except Exception as e:
        status = html.Div(f"An error occurred: {e}", className="text-danger")
        return (no_update, status, {'display': 'none'}, [], {'display': 'none'}, [],
                {'display': 'none'}, [], {'display': 'none'}, [], {'display': 'none'}, 
                [], {'display': 'none'})

# =============================================================================
# Graphing & Simulation Callbacks
# =============================================================================
@app.callback(Output('map-graph', 'figure'),
              [Input('dataframe-store', 'data'), Input('comm-link-selector', 'value'), Input('outage-threshold-input', 'value')])
def update_map_graph(json_data, comm_link, threshold):
    """Updates the main flight path analysis map."""
    if not json_data:
        # Return a basic empty plot when no data
        return go.Figure().update_layout(
            template="plotly_dark", 
            title="Upload flight data to view flight path analysis",
            xaxis=dict(title="Longitude (degrees)"),
            yaxis=dict(title="Latitude (degrees)"),
            annotations=[
                dict(
                    text="No data loaded<br><br>Upload a CSV file to begin analysis",
                    xref="paper", yref="paper",
                    x=0.5, y=0.5, xanchor='center', yanchor='middle',
                    showarrow=False,
                    font=dict(size=16, color="gray")
                )
            ]
        )
    
    try:
        df = pd.read_json(io.StringIO(json_data), orient='split')
        return create_report_map_figure(df, comm_link, threshold)
    except Exception as e:
        print(f"Error creating map: {e}")
        return go.Figure().update_layout(
            template="plotly_dark", 
            title="Error loading flight data",
            annotations=[
                dict(
                    text=f"Error: {str(e)}",
                    xref="paper", yref="paper",
                    x=0.5, y=0.5, xanchor='center', yanchor='middle',
                    showarrow=False,
                    font=dict(size=14, color="red")
                )
            ]
        )


# --- NEW: Simulation Callbacks ---
@app.callback(
    [Output('sim-interval', 'disabled'), Output('play-pause-btn', 'children'), Output('play-pause-btn', 'color')],
    [Input('play-pause-btn', 'n_clicks'), Input('stop-btn', 'n_clicks')],
    [State('sim-interval', 'disabled'), State('time-slider', 'value'), State('time-slider', 'max')]
)
def toggle_simulation(play_click, stop_click, is_disabled, slider_val, slider_max):
    """Controls the play, pause, and stop functionality."""
    ctx = dash.callback_context
    if not ctx.triggered:
        return True, "▶ Play", "success"
    
    button_id = ctx.triggered[0]['prop_id'].split('.')[0]
    
    if button_id == 'play-pause-btn':
        if is_disabled:
            # If at the end, restart before playing
            if slider_val == slider_max:
                 # This logic is now handled in the interval callback
                 pass
            return False, "⏸ Pause", "warning"
        else:
            return True, "▶ Play", "success"
    elif button_id == 'stop-btn':
        return True, "▶ Play", "success"
    
    return no_update, no_update, no_update

@app.callback(Output('time-slider', 'value'),
              [Input('sim-interval', 'n_intervals'), Input('stop-btn', 'n_clicks')],
              [State('time-slider', 'value'), State('time-slider', 'max'), State('speed-selector', 'value')])
def update_slider(n_intervals, stop_click, current_value, max_value, speed):
    """Updates the slider position based on the interval timer or stop button."""
    ctx = dash.callback_context
    if ctx.triggered and ctx.triggered[0]['prop_id'].split('.')[0] == 'stop-btn':
        return 0
    
    if current_value is None or max_value is None:
        return 0
        
    # Only advance if we have intervals and speed
    if n_intervals and speed:
        new_value = current_value + (speed or 1)
        if new_value >= max_value:
            return max_value
        return new_value
    
    return current_value or 0


@app.callback(
    [Output('sim-graph', 'figure'), Output('sim-telemetry-display', 'children')],
    [Input('time-slider', 'value'), Input('dataframe-store', 'data'),
     Input('comm-link-selector', 'value'), Input('outage-threshold-input', 'value')],
    prevent_initial_call=True
)
def update_simulation_view(current_time_idx, json_data, comm_link, threshold):
    """Updates the simulation graph and telemetry display based on the slider."""
    if not json_data:
        return go.Figure().update_layout(template="plotly_dark", title_text="Load data to begin simulation"), "No data loaded"
    
    if current_time_idx is None:
        current_time_idx = 0
        
    df = pd.read_json(io.StringIO(json_data), orient='split')
    df['Timestamp'] = pd.to_datetime(df['Timestamp'])
    
    # Ensure current_time_idx is within bounds
    current_time_idx = min(max(0, int(current_time_idx)), len(df) - 1)
    
    # --- Data for current point in time ---
    current_data = df.iloc[current_time_idx]
    
    # --- Determine link status ---
    is_outage = False
    if comm_link and threshold is not None and comm_link in df.columns:
        try:
            is_outage = current_data[comm_link] < threshold
        except:
            is_outage = False
    marker_color = "#DC3545" if is_outage else "#00FF00"
    
    # --- Build Figure ---
    fig = go.Figure()
    
    try:
        # 1. Full flight path (background)
        fig.add_trace(go.Scattermapbox(
            mode="lines", 
            lon=df['POS_Longitude_deg'], 
            lat=df['POS_Latitude_deg'],
            line=dict(width=2, color='rgba(128,128,128,0.5)'), 
            name="Full Path",
            showlegend=False
        ))
        
        # 2. Traveled path
        if current_time_idx > 0:
            traveled_df = df.iloc[:current_time_idx+1]
            fig.add_trace(go.Scattermapbox(
                mode="lines", 
                lon=traveled_df['POS_Longitude_deg'], 
                lat=traveled_df['POS_Latitude_deg'],
                line=dict(width=4, color='#1F77B4'), 
                name="Traveled Path",
                showlegend=False
            ))
        
        # 3. Current aircraft position
        fig.add_trace(go.Scattermapbox(
            mode="markers", 
            lon=[current_data['POS_Longitude_deg']], 
            lat=[current_data['POS_Latitude_deg']],
            marker=dict(size=15, color=marker_color, symbol='circle'), 
            name="Current Position",
            showlegend=False
        ))
        
        # Try mapbox layout
        fig.update_layout(
            template="plotly_dark", 
            showlegend=False,
            mapbox=dict(
                style="open-street-map",  # Changed to open-street-map for better reliability
                center=dict(
                    lon=current_data['POS_Longitude_deg'],
                    lat=current_data['POS_Latitude_deg']
                ),
                zoom=10
            ),
            margin={"r":0,"t":30,"l":0,"b":0},
            title=f"Flight Simulation - Time Index: {current_time_idx}/{len(df)-1}"
        )
        
    except Exception as e:
        print(f"Mapbox simulation error: {e}")
        # Fallback to basic scatter plot
        fig = go.Figure()
        
        # Convert to regular scatter plot
        fig.add_trace(go.Scatter(
            x=df['POS_Longitude_deg'], 
            y=df['POS_Latitude_deg'],
            mode="lines", 
            line=dict(width=2, color='rgba(128,128,128,0.5)'),
            name="Full Path"
        ))
        
        if current_time_idx > 0:
            traveled_df = df.iloc[:current_time_idx+1]
            fig.add_trace(go.Scatter(
                x=traveled_df['POS_Longitude_deg'], 
                y=traveled_df['POS_Latitude_deg'],
                mode="lines", 
                line=dict(width=4, color='#1F77B4'),
                name="Traveled Path"
            ))
        
        fig.add_trace(go.Scatter(
            x=[current_data['POS_Longitude_deg']], 
            y=[current_data['POS_Latitude_deg']],
            mode="markers", 
            marker=dict(size=15, color=marker_color),
            name="Current Position"
        ))
        
        fig.update_layout(
            template="plotly_dark",
            title=f"Flight Simulation (Basic Plot) - Time Index: {current_time_idx}/{len(df)-1}",
            xaxis_title="Longitude (deg)",
            yaxis_title="Latitude (deg)",
            showlegend=True
        )

    # --- Build Telemetry Display ---
    try:
        telemetry_children = [
            html.H5(f"Timestamp: {current_data['Timestamp'].strftime('%Y-%m-%d %H:%M:%S')}", className="mb-2"),
            dbc.Row([
                dbc.Col([
                    html.Strong("Position:"),
                    html.Br(),
                    f"Lat: {current_data['POS_Latitude_deg']:.4f}°",
                    html.Br(),
                    f"Lon: {current_data['POS_Longitude_deg']:.4f}°"
                ], width=3),
                dbc.Col([
                    html.Strong("Altitude:"),
                    html.Br(),
                    f"{current_data['POS_Altitude_ft']:,.0f} ft"
                ], width=2),
                dbc.Col([
                    html.Strong("Communication:"),
                    html.Br(),
                    html.Span(
                        f"{comm_link.replace('_', ' ')}: {current_data.get(comm_link, 'N/A'):.2f} dB" if comm_link and comm_link in current_data else "No comm link selected",
                        style={'color': marker_color if comm_link else 'white'}
                    )
                ], width=4),
                dbc.Col([
                    html.Strong("Progress:"),
                    html.Br(),
                    f"{(current_time_idx / (len(df)-1) * 100):.1f}% complete"
                ], width=3)
            ])
        ]
    except Exception as e:
        telemetry_children = [html.Div(f"Error displaying telemetry: {str(e)}")]
    
    return fig, telemetry_children


# --- Other Graph Callbacks (Unchanged from original) ---
@app.callback(Output('main-graph', 'figure'),
              Input('subsystem-checklist', 'value'), State('dataframe-store', 'data'))
def update_2d_graph(selected_subsystems, json_data):
    if not selected_subsystems or not json_data:
        return go.Figure().update_layout(template="plotly_dark", title_text="Select a subsystem to see telemetry")
    df = pd.read_json(io.StringIO(json_data), orient='split')
    fig = go.Figure()
    for subsystem in selected_subsystems:
        subsystem_cols = [col for col in df.columns if col.startswith(f"{subsystem}_")]
        for col in subsystem_cols:
            fig.add_trace(go.Scatter(x=df['Timestamp'], y=df[col], mode='lines', name=col))
    return fig.update_layout(title="Subsystem Telemetry Over Time", template="plotly_dark")

@app.callback(
    Output('scatter-3d-graph', 'figure'),
    [Input('x-axis-selector', 'value'), Input('y-axis-selector', 'value'),
     Input('z-axis-selector', 'value'), Input('color-selector', 'value')],
    State('dataframe-store', 'data')
)
def update_3d_scatter(x_axis, y_axis, z_axis, color_axis, json_data):
    if not all([x_axis, y_axis, z_axis, color_axis, json_data]):
        return go.Figure().update_layout(template="plotly_dark", title="Select axes to generate 3D plot")
    df = pd.read_json(io.StringIO(json_data), orient='split')
    fig = go.Figure(data=[go.Scatter3d(
        x=df[x_axis], y=df[y_axis], z=df[z_axis], mode='markers',
        marker=dict(size=4, color=df[color_axis], colorscale='Viridis', showscale=True,
                    colorbar_title_text=color_axis.replace('_', ' '))
    )])
    fig.update_layout(template="plotly_dark", margin=dict(l=0, r=0, b=0, t=40),
                      scene=dict(xaxis_title=x_axis, yaxis_title=y_axis, zaxis_title=z_axis))
    return fig

# =============================================================================
# Report Generation Callbacks (Now with Plots!)
# =============================================================================
@app.callback(
    Output("download-pdf", "data"),
    Input("download-pdf-btn", "n_clicks"),
    [State("dataframe-store", "data"), State("comm-link-selector", "value"), State("outage-threshold-input", "value")],
    prevent_initial_call=True
)
def generate_pdf_report(n_clicks, json_data, comm_link, threshold):
    if not json_data: return no_update
    df = pd.read_json(io.StringIO(json_data), orient="split")
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        doc = SimpleDocTemplate(tmp_file.name)
        styles = getSampleStyleSheet()
        story = [Paragraph("Flight Data Report", styles["Title"])]

        # --- Generate and save plot as an in-memory image ---
        fig = create_report_map_figure(df, comm_link, threshold)
        img_buffer = io.BytesIO()
        fig.write_image(img_buffer, format="png", width=800, height=600, scale=2)
        img_buffer.seek(0)
        
        # --- Add image to PDF ---
        story.append(Spacer(1, 0.2 * inch))
        story.append(Image(img_buffer, width=6*inch, height=4.5*inch))
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("This plot shows the full flight path. Red segments indicate periods where the selected communication link margin dropped below the specified threshold.", styles["Normal"]))

        doc.build(story)
        return dcc.send_file(tmp_file.name, filename="Flight_Report.pdf")

@app.callback(
    Output("download-word", "data"),
    Input("download-word-btn", "n_clicks"),
    [State("dataframe-store", "data"), State("comm-link-selector", "value"), State("outage-threshold-input", "value")],
    prevent_initial_call=True
)
def generate_word_report(n_clicks, json_data, comm_link, threshold):
    if not json_data: return no_update
    df = pd.read_json(io.StringIO(json_data), orient="split")
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        doc = Document()
        doc.add_heading("Flight Data Report", 0)

        # --- Generate and save plot as an in-memory image ---
        fig = create_report_map_figure(df, comm_link, threshold)
        img_buffer = io.BytesIO()
        fig.write_image(img_buffer, format="png", width=800, height=600, scale=2)
        img_buffer.seek(0)
        
        # --- Add image to Word doc ---
        doc.add_picture(img_buffer, width=Inches(6.0))
        doc.add_paragraph("This plot shows the full flight path. Red segments indicate periods where the selected communication link margin dropped below the specified threshold.")
        
        doc.save(tmp_file.name)
        return dcc.send_file(tmp_file.name, filename="Flight_Report.docx")


if __name__ == '__main__':
    app.run(debug=True)